#!/usr/bin/env python3
"""Convert HTML files to Word documents."""

import re
import os
from bs4 import BeautifulSoup
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


def set_cell_shading(cell, color_hex):
    """Set background color for a table cell."""
    shading = OxmlElement('w:shd')
    shading.set(qn('w:fill'), color_hex)
    cell._tc.get_or_add_tcPr().append(shading)


def set_paragraph_shading(paragraph, color_hex):
    """Set background color for a paragraph."""
    shading = OxmlElement('w:shd')
    shading.set(qn('w:fill'), color_hex)
    paragraph._p.get_or_add_pPr().append(shading)


def convert_highlighted_code(html_path, output_path):
    """Convert highlighted_code.html to Word document using BeautifulSoup."""
    with open(html_path, 'r', encoding='utf-8') as f:
        soup = BeautifulSoup(f.read(), 'html.parser')

    doc = Document()

    # Set narrow margins
    for section in doc.sections:
        section.left_margin = Inches(0.5)
        section.right_margin = Inches(0.5)
        section.top_margin = Inches(0.5)
        section.bottom_margin = Inches(0.5)

    # Add legend
    legend = doc.add_paragraph()
    legend.add_run("Legend: ").bold = True
    doc.add_paragraph()

    legend_table = doc.add_table(rows=1, cols=2)
    legend_table.alignment = WD_TABLE_ALIGNMENT.LEFT

    cell1 = legend_table.cell(0, 0)
    cell1.text = "GREEN = Original Code (Written by myself)"
    set_cell_shading(cell1, "90EE90")

    cell2 = legend_table.cell(0, 1)
    cell2.text = "WHITE = Starter/Template Code"
    set_cell_shading(cell2, "FFFFFF")

    doc.add_paragraph()

    # Find all file headers and their following pre blocks
    file_headers = soup.find_all('div', class_='file-header')

    for header in file_headers:
        file_name = header.get_text(strip=True)

        # Add file header
        h = doc.add_paragraph()
        h_run = h.add_run(file_name)
        h_run.bold = True
        h_run.font.size = Pt(11)
        h_run.font.color.rgb = RGBColor(255, 255, 255)
        set_paragraph_shading(h, '333333')
        h.paragraph_format.space_after = Pt(0)

        # Find the next pre element (sibling)
        pre = header.find_next_sibling('pre')
        if not pre:
            continue

        # Find all line spans within this pre
        lines = pre.find_all('span', class_=lambda x: x and 'line' in x and 'line-num' not in x)

        if lines:
            # Create a table with one row per line
            table = doc.add_table(rows=len(lines), cols=1)

            for i, line_span in enumerate(lines):
                classes = line_span.get('class', [])
                is_original = 'original' in classes

                # Get the text content, excluding line numbers
                line_text = ""
                for child in line_span.children:
                    if hasattr(child, 'get') and 'line-num' in child.get('class', []):
                        continue
                    if hasattr(child, 'get_text'):
                        line_text += child.get_text()
                    else:
                        line_text += str(child)

                cell = table.cell(i, 0)
                p = cell.paragraphs[0]
                run = p.add_run(line_text)
                run.font.name = 'Courier New'
                run.font.size = Pt(8)

                # Set background color
                if is_original:
                    set_cell_shading(cell, "90EE90")  # Light green
                else:
                    set_cell_shading(cell, "FFFFFF")  # White

        doc.add_paragraph()

    doc.save(output_path)
    print(f"Saved: {output_path}")


def convert_report(html_path, output_path, base_dir):
    """Convert CM3020_PartB_Report_Final.html to Word document using BeautifulSoup."""
    with open(html_path, 'r', encoding='utf-8') as f:
        soup = BeautifulSoup(f.read(), 'html.parser')

    doc = Document()

    # Set margins
    for section in doc.sections:
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)

    # Get the body content
    body = soup.find('body')
    if not body:
        print("No body found!")
        return

    def process_text_with_formatting(element, paragraph):
        """Process text with inline formatting (bold, code, etc.)"""
        for child in element.children:
            if isinstance(child, str):
                text = child
                if text:
                    paragraph.add_run(text)
            elif child.name == 'strong':
                run = paragraph.add_run(child.get_text())
                run.bold = True
            elif child.name == 'code':
                run = paragraph.add_run(child.get_text())
                run.font.name = 'Courier New'
                run.font.size = Pt(10)
            elif child.name == 'sub':
                run = paragraph.add_run(child.get_text())
                run.font.subscript = True
            else:
                # Recursively process other elements
                process_text_with_formatting(child, paragraph)

    # Process each direct child of body
    for element in body.children:
        if isinstance(element, str):
            continue

        if element.name == 'style':
            continue

        elif element.name == 'h1':
            p = doc.add_heading(element.get_text(strip=True), level=0)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER

        elif element.name == 'h2':
            doc.add_heading(element.get_text(strip=True), level=1)

        elif element.name == 'h3':
            doc.add_heading(element.get_text(strip=True), level=2)

        elif element.name == 'p':
            text = element.get_text(strip=True)
            if not text:
                continue

            p = doc.add_paragraph()
            process_text_with_formatting(element, p)

            # Check for special classes
            classes = element.get('class', [])
            if 'word-count' in classes:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in p.runs:
                    run.bold = True
            else:
                p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

        elif element.name == 'pre':
            p = doc.add_paragraph()
            run = p.add_run(element.get_text())
            run.font.name = 'Courier New'
            run.font.size = Pt(9)
            set_paragraph_shading(p, 'F5F5F5')

        elif element.name == 'table':
            rows = element.find_all('tr')
            if rows:
                # Determine number of columns
                first_row_cells = rows[0].find_all(['th', 'td'])
                num_cols = len(first_row_cells)

                table = doc.add_table(rows=len(rows), cols=num_cols)
                table.style = 'Table Grid'

                for i, row in enumerate(rows):
                    cells = row.find_all(['th', 'td'])
                    for j, cell_elem in enumerate(cells):
                        if j < num_cols:
                            cell = table.cell(i, j)
                            cell.text = cell_elem.get_text(strip=True)

                            if cell_elem.name == 'th':
                                for para in cell.paragraphs:
                                    for run in para.runs:
                                        run.bold = True
                                set_cell_shading(cell, 'F0F0F0')

                doc.add_paragraph()

        elif element.name == 'div':
            classes = element.get('class', [])

            if 'figure' in classes:
                # Find all images in this figure
                images = element.find_all('img')

                # Find screenshot grid if present
                screenshot_grid = element.find('div', class_='screenshot-grid')

                if screenshot_grid:
                    # Handle grid of screenshots
                    screenshot_items = screenshot_grid.find_all('div', class_='screenshot-item')

                    if len(screenshot_items) >= 2:
                        # Create a 2-column table for the images
                        num_rows = (len(screenshot_items) + 1) // 2
                        img_table = doc.add_table(rows=num_rows * 2, cols=2)

                        for idx, item in enumerate(screenshot_items):
                            row_idx = (idx // 2) * 2
                            col_idx = idx % 2

                            img_elem = item.find('img')
                            caption_elem = item.find('div', class_='screenshot-caption')

                            if img_elem:
                                src = img_elem.get('src', '')
                                img_path = os.path.join(base_dir, src)

                                if os.path.exists(img_path):
                                    cell = img_table.cell(row_idx, col_idx)
                                    p = cell.paragraphs[0]
                                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                                    run = p.add_run()
                                    run.add_picture(img_path, width=Inches(2.8))

                            if caption_elem:
                                caption_cell = img_table.cell(row_idx + 1, col_idx)
                                cap_p = caption_cell.paragraphs[0]
                                cap_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                                cap_run = cap_p.add_run(caption_elem.get_text(strip=True))
                                cap_run.italic = True
                                cap_run.font.size = Pt(10)

                elif images:
                    # Single image(s) outside of grid
                    for img_elem in images:
                        src = img_elem.get('src', '')
                        img_path = os.path.join(base_dir, src)

                        if os.path.exists(img_path):
                            p = doc.add_paragraph()
                            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                            run = p.add_run()
                            run.add_picture(img_path, width=Inches(5))

                # Add figure caption
                figure_caption = element.find('div', class_='figure-caption')
                if figure_caption:
                    cap_p = doc.add_paragraph()
                    cap_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    cap_run = cap_p.add_run(figure_caption.get_text(strip=True))
                    cap_run.italic = True
                    cap_run.font.size = Pt(10)

                doc.add_paragraph()

    doc.save(output_path)
    print(f"Saved: {output_path}")


if __name__ == '__main__':
    base_dir = '/Users/cj/Desktop/src'
    desktop = '/Users/cj/Desktop'

    # Convert highlighted code
    print("Converting highlighted_code.html...")
    convert_highlighted_code(
        os.path.join(base_dir, 'highlighted_code.html'),
        os.path.join(desktop, 'highlighted_code.docx')
    )

    # Convert report
    print("Converting CM3020_PartB_Report_Final.html...")
    convert_report(
        os.path.join(base_dir, 'CM3020_PartB_Report_Final.html'),
        os.path.join(desktop, 'CM3020_PartB_Report.docx'),
        base_dir
    )

    print("\nDone! Files saved to Desktop.")
